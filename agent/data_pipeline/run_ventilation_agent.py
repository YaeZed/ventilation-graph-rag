import logging
import os
import json
import dataclasses
from docx import Document
import re
from dotenv import load_dotenv

# 解决作为模块被外部调用时的路径问题
import sys
current_dir = os.path.dirname(os.path.abspath(__file__))
if current_dir not in sys.path:
    sys.path.append(current_dir)

from ventilation_safety_agent import VentilationSafetyAgent
from ventilation_kg_builder import VentilationKGBuilder

logger = logging.getLogger(__name__)
def setup_logging():
    """配置logging（在程序入口调用一次）"""
    os.makedirs('safety_output', exist_ok=True)  # 确保输出目录存在
    logging.basicConfig(
        # 打印info及以上的日志
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S',
        handlers=[
            logging.StreamHandler(),  # 输出到控制台
            logging.FileHandler('safety_output/ventilation.log', encoding='utf-8')  # 输出到文件
        ]
    )

def read_docx_with_tables(file_path:str):
    """
    从Word文档中读取内容，将表格转换为Markdown格式
    
    Args:
        file_path: Word文档路径
    
    Returns:
        str: 转换后的文本内容，表格以Markdown格式呈现
    """
    # 1.函数入口用info
    logger.info(f"开始读取word文档:{file_path}")

    try:
        # 2.主要步骤debug
        logger.debug(f"开始加载文档：{file_path}")
        doc = Document(file_path)
        if doc is None:
            logger.error("Document对象为None")
            return ''

        logger.debug(f"文档加载成功，开始处理\n")

        # 切分的文档内容，段落数，表数
        content = []
        para_count = 0
        table_count = 0

        # 按文档顺序返回所有段落和表格的XML元素
        for idx,element in enumerate(doc.element.body):
            """
            element.tag 的值示例:
            段落: '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'
            表格: '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl'
            """
            if element.tag.endswith("p"):
                # 找到对应的段落对象
                para = [p for p in doc.paragraphs if p._element == element][0]
                # 过滤空段落
                if para.text.strip():
                    content.append(para.text)
                    para_count += 1
                    # 3.详细信息：DEBUG（可选，调试时有用）
                    # logger.debug(f"处理段落 {para_count}: {para.text[:50]}...")

            elif element.tag.endswith("tbl"):
                # 找到对应的表格对象
                table = [t for t in doc.tables if t._element == element][0]
                # 过滤空表格
                if len(table.rows) == 0:
                    logger.warning(f"发现空表格（索引 {idx}），跳过")
                    continue
                table_count += 1
                # logger.debug(f"处理表格 {table_count}，行数: {len(table.rows)}")
                
                md_rows = []
                for idx,row in enumerate(table.rows):
                    cells = [cell.text.strip().replace('\n', ' ').replace('\r', '') for cell in row.cells]
                    md_rows.append("| " + " | ".join(cells) + " |")

                    # 在第一行后添加表头分隔符（标准Markdown格式）
                    if idx == 0:
                        separator = "|" + "|".join([" --- " for _ in cells]) + "|"
                        md_rows.append(separator)

                # 在表格前后添加空行，确保Markdown渲染正确
                content.append("\n" + "\n".join(md_rows) + "\n")
        # 4.函数完成info
        logger.info(f"✅ 文档读取完成，共 {len(content)} 个元素，{para_count}个段落元素，{table_count}个表格元素")
        return "\n".join(content)            

    # 5.异常处理：error+exception    
    except FileNotFoundError:
        logger.error(f"文件不存在: {file_path}")
        raise
    except Exception as e:
        logger.error(f"读取文档时发生错误: {type(e).__name__} - {str(e)}")
        logger.exception("详细错误堆栈:")  # 自动记录完整的异常堆栈
        raise
            
def split_by_article(text:str):
    """
    按条款切分文本
    
    支持的格式：
    - 第一条、第二条...第九条
    - 第十条、第二十条...第九十条
    - 第一百条、第一百零一条...第九百九十九条
    - 第一百五十六条、第一百五十七条等
    
    Args:
        text: 完整文本
    
    Returns:
        list: 条款列表，每个元素是完整的条款内容（含编号）
    """
    # 函数入口：INFO（记录输入规模，方便排查问题）
    logger.info(f"开始切分文本，文本长度: {len(text)} 字符")
    # 正则表达式说明：
    # (?:^|\n)       - 非捕获组：匹配行首（^）或换行符（\n），确保条款编号只在行首匹配
    # (第...条\s*)   - 捕获组：匹配条款编号本身（供 re.split 保留）
    # re.MULTILINE   - 让 ^ 匹配每一行的行首，而不只是整个字符串的开头
    pattern = r'(?:^|\n)(第[一二三四五六七八九百十零]+条\s*)'

    parts = re.split(pattern,text)
    articles = []

    # re.split 会将匹配的内容也放入结果列表
    # 奇数索引是条款编号，偶数索引是条款内容
    for i in range(1,len(parts),2):
        if i + 1 < len(parts):
            article_text = parts[i] + parts[i + 1]
            # 过滤掉过短的条款（可能是误匹配）
            if len(article_text.strip()) > 10:
                articles.append(article_text.strip())
    # 2.函数完成：INFO（记录输出规模）
    logger.info(f"✅ 切分完成，共 {len(articles)} 个条款")
    return articles

def extract_article_id(article_text:str):
    """
    从条款文本中提取条款编号
    
    Args:
        article_text: 条款完整文本
    
    Returns:
        str: 条款编号（如"156"），如果未找到则返回"unknown"
    """
    
    match = re.search(r'第([一二三四五六七八九百十零]+)条', article_text)
    if match:
        chinese_num = match.group(1)
        # 转换中文数字为阿拉伯数字
        try:
            arabic_num = chinese_to_arabic(chinese_num)
            return str(arabic_num)
        except:
            return chinese_num  # 如果转换失败，返回中文数字
    return "unknown"

def chinese_to_arabic(chinese_num:str):
    """
    将中文数字转换为阿拉伯数字
    
    Args:
        chinese_num: 中文数字字符串（如"一百五十六"）
    
    Returns:
        int: 阿拉伯数字（如156）
    """
    chinese_map = {
        '零': 0, '一': 1, '二': 2, '三': 3, '四': 4,
        '五': 5, '六': 6, '七': 7, '八': 8, '九': 9,
        '十': 10, '百': 100, '千': 1000
    }
    
    result = 0
    temp = 0
    
    for char in chinese_num:
        if char in ['十', '百', '千']:
            if temp == 0:
                temp = 1
            result += temp * chinese_map[char]
            temp = 0
        else:
            temp = chinese_map.get(char, 0)
    
    result += temp
    return result

def test_agent_json(articles: list, num: int = 3):
    """
    测试Agent处理结果，将前num个条款的JSON输出保存到文件
    
    Args:
        articles: 条款列表
        num: 要测试的条款数量（默认3个，避免API调用过多）
    """
    import time
    
    api_key = os.getenv("DASHSCOPE_API_KEY")
    base_url = "https://dashscope.aliyuncs.com/compatible-mode/v1"
    
    if not api_key:
        logger.error("未找到DASHSCOPE_API_KEY环境变量，请先配置API密钥")
        print("❌ 请先在.env中配置 DASHSCOPE_API_KEY")
        return
    
    agent = VentilationSafetyAgent(api_key, base_url)
    results = []
    
    print(f"\n{'='*60}")
    print(f"🤖 调用LLM解析前{num}个条款...")
    print(f"{'='*60}")
    
    for idx, art_text in enumerate(articles[:num], 1):
        print(f"\n[{idx}/{num}] 正在解析: {art_text[:20].strip()}...")
        
        parsed = agent.extract_logic(art_text)
        
        if parsed:
            # dataclasses.asdict() 将dataclass对象递归转换为字典
            result_dict = dataclasses.asdict(parsed)
            results.append(result_dict)
            print(f"  ✅ 解析成功: 指标数={len(parsed.metrics)}, 要求数={len(parsed.requirements)}")
        else:
            print(f"  ❌ 解析失败")
        
        # 控制API调用频率（避免超限）
        if idx < num:
            time.sleep(1)
    
    # 保存为格式化的JSON文件
    output_path = "safety_output/parsed_articles.json"
    with open(output_path, "w", encoding="utf-8") as f:
        # indent=2 让JSON有缩进，ensure_ascii=False 保留中文
        json.dump(results, f, ensure_ascii=False, indent=2)
    
    print(f"\n💾 JSON结果已保存到: {output_path}")
    print(f"📊 成功解析 {len(results)}/{num} 个条款")

def main():
    """主程序"""
    setup_logging()
    # 加载 .env 文件中的环境变量（让 os.getenv 能读取到 DASHSCOPE_API_KEY）
    load_dotenv(dotenv_path="../../../.env")
    print("=" * 60)
    print("📚 通风规程知识库构建 - 文档转换测试")
    print("=" * 60)
    # 步骤1: 读取Word文档并转换表格为Markdown
    raw_text = read_docx_with_tables("通风.docx")

    # 保存原始转换结果（用于调试）
    with open("safety_output/raw_text_with_markdown.txt", "w", encoding="utf-8") as f:
        f.write(raw_text)
    print(f"💾 原始文本已保存到: safety_output/raw_text_with_markdown.txt\n")

    # 步骤2: 按条款切分
    articles = split_by_article(raw_text)
    test_agent_json(articles,num=3)
    # 步骤3: 预览前3个条款
    print("\n" + "=" * 60)
    print("📋 条款预览（前3个）:")
    print("=" * 60)
    for idx, article in enumerate(articles[:3], 1):
        article_id = extract_article_id(article)
        print(f"\n【条款 {idx}】ID: {article_id}")
        print("-" * 60)
        # 只显示前200字符
        preview = article[:200] + "..." if len(article) > 200 else article
        print(preview)
        print()

    # 步骤4: 统计信息
    print("\n" + "=" * 60)
    print("📊 文档统计:")
    print("=" * 60)
    print(f"✓ 总条款数: {len(articles)}")
    print(f"✓ 平均条款长度: {sum(len(a) for a in articles) // len(articles) if articles else 0} 字符")
    print(f"✓ 最长条款: {max((len(a) for a in articles), default=0)} 字符")
    print(f"✓ 最短条款: {min((len(a) for a in articles), default=0)} 字符")

    # 保存所有条款（用于后续处理）
    with open("safety_output/articles.txt", "w", encoding="utf-8") as f:
        for article in articles:
            article_id = extract_article_id(article)
            f.write(f"===== 条款ID: {article_id} =====\n")
            f.write(article)
            f.write("\n\n" + "=" * 80 + "\n\n")
    print(f"\n💾 所有条款已保存到: safety_output/articles.txt")

    print("\n" + "=" * 60)
    print("✅ 文档转换测试完成！")
    print("=" * 60)

    # ============================================================
    # 步骤5: 完整流程 - LLM解析 + 图谱构建 + 导出CSV
    # ============================================================
    print("\n" + "=" * 60)
    print("🤖 步骤5: 启动完整知识图谱构建流程")
    print("=" * 60)
    
    api_key = os.getenv("DASHSCOPE_API_KEY")
    if not api_key:
        logger.error("未找到DASHSCOPE_API_KEY，跳过图谱构建步骤")
        print("❌ 请先配置 DASHSCOPE_API_KEY 环境变量")
    else:
        import time
        base_url = "https://dashscope.aliyuncs.com/compatible-mode/v1"
        
        agent = VentilationSafetyAgent(api_key, base_url)
        builder = VentilationKGBuilder()
        
        success_count = 0
        fail_count = 0
        
        logger.info(f"开始处理全部 {len(articles)} 个条款")
        
        for idx, art_text in enumerate(articles, 1):
            logger.info(f"处理进度: [{idx}/{len(articles)}] {art_text[:15].strip()}...")
            print(f"  [{idx}/{len(articles)}] {art_text[:20].strip()}...")
            
            parsed = agent.extract_logic(art_text)
            
            if parsed:
                builder.process_article(parsed)
                success_count += 1
            else:
                fail_count += 1
                logger.warning(f"条款解析失败，跳过: {art_text[:20].strip()}")
            
            # 控制API调用频率，避免触发限流
            time.sleep(1)
        
        # 导出 Neo4j CSV
        print(f"\n✅ 解析完成，成功: {success_count}，失败: {fail_count}")
        builder.export_to_neo4j_csv()
        
        logger.info(f"完整流程结束，成功: {success_count}，失败: {fail_count}")
    
    print("\n" + "=" * 60)
    print("✅ 全部流程完成！")
    print("=" * 60)
    print("\n📁 输出文件清单:")
    print("  - safety_output/raw_text_with_markdown.txt  (原始转换文本)")
    print("  - safety_output/articles.txt               (切分后的条款)")
    print("  - safety_output/nodes.csv                  (Neo4j节点)")
    print("  - safety_output/relationships.csv          (Neo4j关系)")
    print("  - safety_output/ventilation.log            (运行日志)")






if __name__ == "__main__":
    main()
